from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets" / "workbuddy-user-guide"
OUTPUT = ROOT / "WorkBuddy招聘流水线用户使用指南.docx"

VIOLET = "5A57DC"
VIOLET_DARK = "343247"
VIOLET_LIGHT = "EEEDFF"
GREEN = "168B60"
GREEN_LIGHT = "EAF7F1"
AMBER = "B97922"
AMBER_LIGHT = "FBF0DF"
INK = "17202A"
MUTED = "747B86"
LINE = "E4E5E6"
PAPER = "F7F6F2"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=INK, bold=None, italic=None, font="Arial Unicode MS"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    set_run_font(run, size=8.5, color=VIOLET, bold=True)
    run.font.letter_spacing = Pt(0.8)
    keep_with_next(p)
    return p


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    set_run_font(run, size=29, color=VIOLET_DARK, bold=True)
    if subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(0)
        sub.paragraph_format.space_after = Pt(15)
        r = sub.add_run(subtitle)
        set_run_font(r, size=12, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color=VIOLET if level < 3 else VIOLET_DARK,
        bold=True,
    )
    return p


def add_body(doc, text, bold_lead=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=10.5, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5, color=INK)
    return p


def add_callout(doc, label, text, tone="violet"):
    fill, accent = {
        "violet": (VIOLET_LIGHT, VIOLET),
        "green": (GREEN_LIGHT, GREEN),
        "amber": (AMBER_LIGHT, AMBER),
    }[tone]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=9.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_step(doc, number, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    badge = p.add_run(f"{number:02d}  ")
    set_run_font(badge, size=10, color=VIOLET, bold=True)
    heading = p.add_run(title)
    set_run_font(heading, size=10.5, color=INK, bold=True)
    detail = doc.add_paragraph()
    detail.paragraph_format.left_indent = Inches(0.35)
    detail.paragraph_format.space_before = Pt(0)
    detail.paragraph_format.space_after = Pt(3)
    detail.paragraph_format.line_spacing = 1.2
    run = detail.add_run(text)
    set_run_font(run, size=9.25, color=MUTED)


def add_figure(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(ASSETS / filename), width=Inches(6.0))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(0)
    c.paragraph_format.space_after = Pt(9)
    text = c.add_run(caption)
    set_run_font(text, size=8, color=MUTED, italic=True)


def add_module_grid(doc):
    data = [
        ("岗位标准", "统一门槛、权重与考察维度", "简历筛选", "批量入库、解析与评分"),
        ("候选人", "查看台账、画像与流程轨迹", "面试管理", "排期、资料包与提醒"),
        ("人才池", "沉淀可复用的优质人才", "数据报表", "观察渠道、转化与周期"),
    ]
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [4680, 4680])
    remove_table_borders(table)
    for row, values in zip(table.rows, data):
        for index, cell in enumerate(row.cells):
            set_cell_fill(cell, PAPER if index == 0 else VIOLET_LIGHT)
            title, detail = values[index * 2], values[index * 2 + 1]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(title)
            set_run_font(r, size=10, color=VIOLET_DARK, bold=True)
            d = cell.add_paragraph()
            d.paragraph_format.space_after = Pt(0)
            r2 = d.add_run(detail)
            set_run_font(r2, size=8.5, color=MUTED)
        spacer = OxmlElement("w:trPr")
        row._tr.insert(0, spacer)


def add_status_flow(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    remove_table_borders(table)
    for index, (cell, title, detail) in enumerate(zip(
        table.rows[0].cells,
        ["完成初筛", "人工复筛", "确认面试", "资料就绪"],
        ["自动评分分层", "确认是否推进", "锁定时间与面试官", "画像、题库、评价表"],
    )):
        set_cell_fill(cell, [PAPER, VIOLET_LIGHT, GREEN_LIGHT, AMBER_LIGHT][index])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_run_font(r, size=9, color=INK, bold=True)
        d = cell.add_paragraph()
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        d.paragraph_format.space_after = Pt(0)
        r2 = d.add_run(detail)
        set_run_font(r2, size=7.5, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), VIOLET)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("WorkBuddy 招聘流水线 · 用户使用指南   ")
    set_run_font(r, size=8, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_field = OxmlElement("w:r")
    r_field.extend([fld_begin, instr, fld_end])
    p._p.append(r_field)


def add_page_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("WORKBUDDY  /  RECRUITING PIPELINE")
    set_run_font(r, size=7.5, color=MUTED, bold=True)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after, color in (
        (1, 16, 18, 10, VIOLET),
        (2, 13, 14, 7, VIOLET),
        (3, 12, 10, 5, VIOLET_DARK),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    configure_styles(doc)
    add_page_header(section)
    add_footer(section)

    # Cover — compact_reference_guide preset + editorial_cover opening pattern.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(30)
    add_kicker(doc, "WorkBuddy · Recruiting Operations")
    add_title(
        doc,
        "招聘流水线用户使用指南",
        "面向 HR、招聘负责人和用人经理的图文操作手册",
    )
    add_callout(
        doc,
        "一句话介绍",
        "把岗位标准、简历入库、候选人推进、面试筹备和招聘数据统一到一个工作空间。",
        "violet",
    )
    add_figure(doc, "01-overview.png", "工作台总览：从待办、招聘漏斗到岗位进度一屏掌握")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用版本：在线私有版  |  2026 年 7 月")
    set_run_font(r, size=8.5, color=MUTED)

    doc.add_page_break()
    add_kicker(doc, "01 · PRODUCT TOUR")
    add_title(doc, "先用 3 分钟认识网站", "登录后，左侧是功能入口，顶部是全局操作区，中间是当前模块的工作内容。")
    add_heading(doc, "网站能帮你完成什么", 2)
    add_body(
        doc,
        "WorkBuddy 将招聘过程组织为一条连续流水线。HR 负责简历入库、邀约、排期和台账维护；用人经理负责复筛、查看画像与面试评价。双方看到的是同一份候选人状态。",
    )
    add_module_grid(doc)
    add_heading(doc, "推荐的日常工作顺序", 2)
    add_status_flow(doc)
    add_callout(
        doc,
        "登录说明",
        "网站采用私有访问方式。打开访问地址后，按照页面提示完成身份验证即可进入。",
        "green",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("访问地址：")
    set_run_font(r, size=9.5, color=INK, bold=True)
    add_hyperlink(
        p,
        "workbuddy-hiring-pipeline.wangshengxiang7.chatgpt.site",
        "https://workbuddy-hiring-pipeline.wangshengxiang7.chatgpt.site",
    )

    doc.add_page_break()
    add_kicker(doc, "02 · RESUME INTAKE")
    add_title(doc, "批量上传并处理简历", "支持 PDF、Word 和常见图片格式，上传后进入私有存储与处理队列。")
    add_figure(doc, "02-resume-screening.png", "简历筛选页：左侧展示处理步骤，右侧展示批次队列与结果")
    add_step(doc, 1, "进入“简历筛选”", "点击左侧导航“简历筛选”，查看当前批次和自动扫描状态。")
    add_step(doc, 2, "选择文件", "点击页面右上角“上传简历文件”，或在队列底部点击“浏览文件”。")
    add_step(doc, 3, "等待安全入库", "文件上传后，系统记录批次、格式、大小、上传人和处理状态；刷新页面后队列仍会保留。")
    add_step(doc, 4, "查看处理结果", "在队列中查看目标岗位、处理状态、匹配分数和筛选结论。")
    add_callout(
        doc,
        "文件要求",
        "支持 PDF、DOC、DOCX、PNG、JPG、JPEG、WebP；单份不超过 15MB，单批最多 100 份。",
        "amber",
    )

    doc.add_page_break()
    add_kicker(doc, "03 · CANDIDATE LEDGER")
    add_title(doc, "在候选人台账中持续推进", "所有候选人的岗位、分数、状态、负责人和更新时间集中展示。")
    add_figure(doc, "03-candidate-ledger.png", "候选人总台账：可按状态筛选，并查看每位候选人的当前进度")
    add_step(doc, 1, "使用状态筛选", "点击“强推荐”“待复筛”“已邀约”等状态，快速定位当前需要处理的人。")
    add_step(doc, 2, "搜索候选人", "可按姓名、公司或技能搜索；高级筛选将用于岗位、渠道和负责人组合筛选。")
    add_step(doc, 3, "打开候选人详情", "点击候选人所在行，右侧会打开候选人画像和下一步操作。")
    add_callout(
        doc,
        "数据同步",
        "顶部显示“数据已同步”时，候选人状态已保存到正式台账；刷新页面或更换设备后仍可保留。",
        "green",
    )

    doc.add_page_break()
    add_kicker(doc, "04 · CANDIDATE PROFILE")
    add_title(doc, "读懂候选人画像与风险提示", "详情侧栏把简历信息转化为面试官可以快速使用的判断材料。")
    add_figure(doc, "04-candidate-profile.png", "候选人画像：匹配度、核心亮点、风险方向和流程轨迹集中呈现")
    add_step(doc, 1, "先看匹配度与状态", "匹配度反映候选人与岗位标准的综合适配程度，状态决定当前需要执行的动作。")
    add_step(doc, 2, "核对匹配点与风险", "核心匹配点用于判断优势，风险与深挖方向用于设计面试追问。")
    add_step(doc, 3, "执行下一步", "点击底部主按钮，可以依次完成复筛、保存邀约、确认面试和生成资料包。")
    add_status_flow(doc)
    add_callout(
        doc,
        "操作记录",
        "每次状态变更都会写入流程轨迹和操作日志，方便招聘负责人回溯。",
        "violet",
    )

    doc.add_page_break()
    add_kicker(doc, "05 · INTERVIEW OPERATIONS")
    add_title(doc, "安排面试并检查资料包", "面试管理页同时展示日程、面试形式、面试官和资料准备状态。")
    add_figure(doc, "05-interview.png", "面试管理：按时间查看当天安排，并检查画像、题库和评价表完成度")
    add_step(doc, 1, "确认面试安排", "选择日期、开始时间、面试形式和面试官，确认后候选人进入“已确认面试”。")
    add_step(doc, 2, "检查资料准备状态", "右侧完成度区域会提示一页纸画像、定制题库和评价表是否准备完成。")
    add_step(doc, 3, "打开资料包", "切换到“资料包”视图，查看每位候选人的面试材料。")
    add_callout(
        doc,
        "提醒机制",
        "系统设计为在面试前固定时间提醒 HR 与面试官，并附带面试信息和资料包路径。",
        "amber",
    )

    doc.add_page_break()
    add_kicker(doc, "06 · RECRUITING ANALYTICS")
    add_title(doc, "用招聘报表发现问题", "从简历数量、阶段转化、渠道质量和招聘周期四个角度观察招聘效率。")
    add_figure(doc, "06-reports.png", "招聘数据报表：关键指标、趋势、渠道质量和流程洞察")
    add_step(doc, 1, "先看关键指标", "关注本月简历量、初筛通过率、邀约到场率、平均招聘周期和 AI 判断一致性。")
    add_step(doc, 2, "再看趋势与渠道", "比较近 8 周简历量与复筛通过情况，识别高质量简历渠道。")
    add_step(doc, 3, "处理流程洞察", "对等待时间过长、岗位复筛积压和到场率下降等问题及时催办。")
    add_callout(
        doc,
        "使用建议",
        "每周固定查看一次报表，并结合人工复筛结果校准岗位标准和 AI 评分规则。",
        "green",
    )

    doc.add_page_break()
    add_kicker(doc, "07 · DAILY CHECKLIST")
    add_title(doc, "日常使用建议与常见问题", "按照固定节奏使用，可以减少遗漏并保持招聘数据一致。")
    add_heading(doc, "建议的每日工作节奏", 2)
    add_step(doc, 1, "上午检查新增简历", "确认上传批次是否完成入库，并处理格式错误或重复投递。")
    add_step(doc, 2, "中午完成复筛推进", "优先处理强推荐和等待超过 24 小时的候选人。")
    add_step(doc, 3, "下午确认面试筹备", "检查次日面试安排、面试官和资料包完整性。")
    add_step(doc, 4, "下班前核对台账", "确认当天候选人状态、负责人和备注已经更新。")
    add_heading(doc, "常见问题", 2)
    add_body(doc, "为什么页面显示“演示数据”？", bold_lead="为什么页面显示“演示数据”？")
    add_body(doc, "本地预览或数据服务暂时不可用时，页面会自动使用演示数据。正式站点显示“数据已同步”时，状态会持久保存。")
    add_body(doc, "上传失败怎么办？", bold_lead="上传失败怎么办？")
    add_body(doc, "先检查文件格式、文件大小和单批数量。上传失败时系统会自动清理未完成文件，可修正后重新上传。")
    add_body(doc, "用人经理可以看到什么？", bold_lead="用人经理可以看到什么？")
    add_body(doc, "用人经理主要查看候选人画像、岗位匹配点、风险提示、面试资料和本部门招聘进度。")
    add_callout(
        doc,
        "当前版本说明",
        "真实文件入库、候选人台账和状态持久化已经可用；简历文本提取、OCR、自动去重与 AI 评分仍在继续建设。",
        "violet",
    )
    core = doc.core_properties
    core.title = "WorkBuddy 招聘流水线用户使用指南"
    core.subject = "面向 HR、招聘负责人和用人经理的图文操作手册"
    core.author = "WorkBuddy"
    core.keywords = "招聘流水线, 简历筛选, 候选人管理, 面试管理, 用户指南"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
